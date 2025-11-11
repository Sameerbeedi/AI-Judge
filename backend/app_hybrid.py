from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from sqlalchemy.orm import Session
import os
from dotenv import load_dotenv
from groq import Groq
import PyPDF2
import docx
import io
from datetime import datetime

# Import our database services
from database import get_db, init_db, Case, Document, Argument, Verdict, FollowUp
from neo4j_service import neo4j_service
from cache import cache, cache_verdict, get_cached_verdict, cache_case_data, get_cached_case, invalidate_case_cache

# Load environment variables
load_dotenv()

app = FastAPI(title="AI Judge API", version="2.0.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

# Pydantic models
class CaseCreate(BaseModel):
    case_id: str

class ArgumentSubmit(BaseModel):
    case_id: str
    side: str  # "A" or "B"
    text: str
    
class FollowUpArgument(BaseModel):
    case_id: str
    side: str
    argument: str

# Startup/Shutdown events
@app.on_event("startup")
async def startup():
    """Initialize databases on startup"""
    init_db()  # Create PostgreSQL tables
    print("✅ PostgreSQL initialized")
    
    if neo4j_service.check_connection():
        print("✅ Neo4j connected")
    else:
        print("⚠️  Neo4j not available (optional)")
    
    if cache.check_connection():
        print("✅ Redis connected")
    else:
        print("⚠️  Redis not available (caching disabled)")

@app.on_event("shutdown")
async def shutdown():
    """Close database connections"""
    neo4j_service.close()
    print("🔌 Connections closed")

# Helper functions
def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from various file formats"""
    try:
        if filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        elif filename.endswith('.docx'):
            doc = docx.Document(io.BytesIO(file_content))
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif filename.endswith('.txt'):
            return file_content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {filename}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting text: {str(e)}")

def generate_summary(text: str, max_length: int = 200) -> str:
    """Generate a brief summary for graph storage"""
    words = text.split()
    if len(words) <= max_length:
        return text
    return " ".join(words[:max_length]) + "..."

# API Endpoints
@app.get("/")
def read_root():
    return {
        "message": "AI Judge API - Hybrid Database Version",
        "version": "2.0.0",
        "databases": {
            "postgresql": "Case data & documents",
            "neo4j": "Relationships & patterns",
            "redis": "Caching & performance"
        }
    }

@app.get("/health")
def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "databases": {
            "postgresql": "connected",
            "neo4j": "connected" if neo4j_service.check_connection() else "optional",
            "redis": "connected" if cache.check_connection() else "disabled"
        }
    }

@app.post("/api/case/create")
def create_case(case: CaseCreate, db: Session = Depends(get_db)):
    """Create a new case"""
    # Check if case already exists
    existing_case = db.query(Case).filter(Case.case_id == case.case_id).first()
    if existing_case:
        raise HTTPException(status_code=400, detail="Case ID already exists")
    
    # Create in PostgreSQL
    new_case = Case(case_id=case.case_id, status="pending")
    db.add(new_case)
    db.commit()
    db.refresh(new_case)
    
    # Create in Neo4j
    try:
        neo4j_service.create_case_node(case.case_id)
    except Exception as e:
        print(f"Neo4j creation failed (optional): {e}")
    
    return {
        "message": "Case created successfully",
        "case_id": case.case_id,
        "status": "pending"
    }

@app.post("/api/case/{case_id}/upload")
async def upload_documents(
    case_id: str,
    side: str,
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    """Upload documents for a case"""
    # Verify case exists
    case = db.query(Case).filter(Case.case_id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    if side not in ["A", "B"]:
        raise HTTPException(status_code=400, detail="Side must be 'A' or 'B'")
    
    uploaded_files = []
    
    for file in files:
        content = await file.read()
        extracted_text = extract_text_from_file(content, file.filename)
        
        # Save to PostgreSQL
        document = Document(
            case_id=case_id,
            side=side,
            filename=file.filename,
            extracted_text=extracted_text
        )
        db.add(document)
        uploaded_files.append(file.filename)
    
    db.commit()
    
    # Invalidate cache
    invalidate_case_cache(case_id)
    
    return {
        "message": f"{len(uploaded_files)} documents uploaded for side {side}",
        "files": uploaded_files
    }

@app.post("/api/case/{case_id}/argument")
def submit_argument(
    case_id: str,
    argument: ArgumentSubmit,
    db: Session = Depends(get_db)
):
    """Submit argument text for a side"""
    # Verify case exists
    case = db.query(Case).filter(Case.case_id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    if argument.side not in ["A", "B"]:
        raise HTTPException(status_code=400, detail="Side must be 'A' or 'B'")
    
    # Save to PostgreSQL
    new_argument = Argument(
        case_id=case_id,
        side=argument.side,
        argument_text=argument.text
    )
    db.add(new_argument)
    db.commit()
    
    # Invalidate cache
    invalidate_case_cache(case_id)
    
    return {
        "message": f"Argument submitted for side {argument.side}",
        "case_id": case_id
    }

@app.post("/api/case/{case_id}/verdict")
def generate_verdict(case_id: str, db: Session = Depends(get_db)):
    """Generate AI verdict for a case"""
    # Check cache first
    cached = get_cached_verdict(case_id)
    if cached:
        return {"verdict": cached, "cached": True}
    
    # Get case from PostgreSQL
    case = db.query(Case).filter(Case.case_id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get arguments for both sides
    side_a_args = db.query(Argument).filter(
        Argument.case_id == case_id,
        Argument.side == "A"
    ).all()
    side_b_args = db.query(Argument).filter(
        Argument.case_id == case_id,
        Argument.side == "B"
    ).all()
    
    # Get documents
    side_a_docs = db.query(Document).filter(
        Document.case_id == case_id,
        Document.side == "A"
    ).all()
    side_b_docs = db.query(Document).filter(
        Document.case_id == case_id,
        Document.side == "B"
    ).all()
    
    if not (side_a_args or side_a_docs) or not (side_b_args or side_b_docs):
        raise HTTPException(status_code=400, detail="Both sides must submit arguments or documents")
    
    # Combine texts
    side_a_text = "\n\n".join([arg.argument_text for arg in side_a_args] + 
                              [doc.extracted_text for doc in side_a_docs if doc.extracted_text])
    side_b_text = "\n\n".join([arg.argument_text for arg in side_b_args] + 
                              [doc.extracted_text for doc in side_b_docs if doc.extracted_text])
    
    # Generate verdict using Groq
    prompt = f"""You are an experienced AI Judge with deep knowledge of legal systems and judicial reasoning.

CASE ID: {case_id}

SIDE A ARGUMENTS:
{side_a_text[:2000]}

SIDE B ARGUMENTS:
{side_b_text[:2000]}

Analyze both sides carefully and provide your verdict. Consider:
1. Strength of legal arguments
2. Evidence presented
3. Logical consistency
4. Applicable legal principles

Provide your verdict in the following format:
1. VERDICT: [In favor of Side A / Side B / Split Decision]
2. REASONING: [Detailed legal reasoning]
3. KEY POINTS: [List key legal principles applied]
4. COUNTERARGUMENTS CONSIDERED: [What arguments you weighed]

Be thorough, fair, and base your decision on legal principles and evidence presented."""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": "You are an experienced AI Judge with deep knowledge of legal systems, precedents, and judicial reasoning. You provide fair, balanced, and well-reasoned verdicts."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=2000
    )
    
    verdict_text = response.choices[0].message.content
    
    # Determine winning side for Neo4j
    winning_side = "A" if "favor of Side A" in verdict_text else "B" if "favor of Side B" in verdict_text else "Split"
    
    # Save to PostgreSQL
    verdict = Verdict(
        case_id=case_id,
        verdict_text=verdict_text,
        reasoning=verdict_text
    )
    db.add(verdict)
    case.status = "decided"
    db.commit()
    
    # Save to Neo4j
    try:
        neo4j_service.create_argument_nodes(
            case_id,
            generate_summary(side_a_text),
            generate_summary(side_b_text)
        )
        neo4j_service.create_verdict_node(
            case_id,
            generate_summary(verdict_text),
            winning_side
        )
    except Exception as e:
        print(f"Neo4j save failed (optional): {e}")
    
    # Cache the verdict
    verdict_data = {"verdict": verdict_text, "decided_at": datetime.now().isoformat()}
    cache_verdict(case_id, verdict_data)
    
    return {
        "verdict": verdict_text,
        "case_id": case_id,
        "cached": False
    }

@app.post("/api/case/{case_id}/followup")
def submit_followup(
    case_id: str,
    followup: FollowUpArgument,
    db: Session = Depends(get_db)
):
    """Submit follow-up argument after initial verdict"""
    # Get case and verdict
    case = db.query(Case).filter(Case.case_id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    verdict = db.query(Verdict).filter(Verdict.case_id == case_id).first()
    if not verdict:
        raise HTTPException(status_code=400, detail="No verdict exists for this case")
    
    # Check follow-up limit
    followup_count = db.query(FollowUp).filter(FollowUp.case_id == case_id).count()
    if followup_count >= 5:
        raise HTTPException(status_code=400, detail="Maximum 5 follow-ups reached")
    
    # Get all previous follow-ups
    previous_followups = db.query(FollowUp).filter(FollowUp.case_id == case_id).all()
    
    # Build conversation history
    history = []
    history.append({
        "role": "system",
        "content": "You are an AI Judge. You've already given a verdict. Now listen to follow-up arguments and reconsider if the new points are compelling. Be open to changing your verdict if new evidence or legal reasoning is presented."
    })
    
    # Add initial verdict
    history.append({
        "role": "user",
        "content": f"YOUR INITIAL VERDICT:\n{verdict.verdict_text}"
    })
    
    # Add previous follow-ups
    for prev in previous_followups:
        history.append({
            "role": "user",
            "content": f"Side {prev.side} argues: {prev.argument_text}"
        })
        history.append({
            "role": "assistant",
            "content": prev.response_text
        })
    
    # Add current argument
    history.append({
        "role": "user",
        "content": f"Side {followup.side} now argues: {followup.argument}\n\nPlease reconsider your verdict. If this argument changes your perspective, explain why and issue a revised verdict. If not, explain why the original verdict still stands."
    })
    
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=history,
        temperature=0.4,
        max_tokens=1500
    )
    
    response_text = response.choices[0].message.content
    
    # Save to PostgreSQL
    new_followup = FollowUp(
        case_id=case_id,
        side=followup.side,
        argument_text=followup.argument,
        response_text=response_text
    )
    db.add(new_followup)
    db.commit()
    
    # Save to Neo4j
    try:
        changed_verdict = "revised verdict" in response_text.lower() or "change my decision" in response_text.lower()
        neo4j_service.add_follow_up(
            case_id,
            followup.side,
            generate_summary(followup.argument),
            changed_verdict
        )
    except Exception as e:
        print(f"Neo4j followup save failed (optional): {e}")
    
    # Invalidate cache
    invalidate_case_cache(case_id)
    
    return {
        "response": response_text,
        "followup_count": followup_count + 1,
        "remaining": 5 - (followup_count + 1)
    }

@app.get("/api/case/{case_id}")
def get_case(case_id: str, db: Session = Depends(get_db)):
    """Get complete case information"""
    # Check cache first
    cached = get_cached_case(case_id)
    if cached:
        return {"case": cached, "cached": True}
    
    # Get from PostgreSQL
    case = db.query(Case).filter(Case.case_id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Build response
    case_data = {
        "case_id": case.case_id,
        "status": case.status,
        "created_at": case.created_at.isoformat(),
        "arguments": {
            "side_a": [{"text": arg.argument_text, "created_at": arg.created_at.isoformat()} 
                       for arg in case.arguments if arg.side == "A"],
            "side_b": [{"text": arg.argument_text, "created_at": arg.created_at.isoformat()} 
                       for arg in case.arguments if arg.side == "B"]
        },
        "documents": {
            "side_a": [{"filename": doc.filename, "uploaded_at": doc.uploaded_at.isoformat()} 
                       for doc in case.documents if doc.side == "A"],
            "side_b": [{"filename": doc.filename, "uploaded_at": doc.uploaded_at.isoformat()} 
                       for doc in case.documents if doc.side == "B"]
        },
        "verdict": {
            "text": case.verdict.verdict_text,
            "decided_at": case.verdict.decided_at.isoformat()
        } if case.verdict else None,
        "follow_ups": [
            {
                "side": f.side,
                "argument": f.argument_text,
                "response": f.response_text,
                "created_at": f.created_at.isoformat()
            } for f in case.follow_ups
        ]
    }
    
    # Cache it
    cache_case_data(case_id, case_data)
    
    return {"case": case_data, "cached": False}

@app.get("/api/statistics")
def get_statistics(db: Session = Depends(get_db)):
    """Get overall statistics from all databases"""
    # PostgreSQL stats
    total_cases = db.query(Case).count()
    decided_cases = db.query(Verdict).count()
    total_documents = db.query(Document).count()
    total_followups = db.query(FollowUp).count()
    
    stats = {
        "postgresql": {
            "total_cases": total_cases,
            "decided_cases": decided_cases,
            "pending_cases": total_cases - decided_cases,
            "total_documents": total_documents,
            "total_followups": total_followups
        }
    }
    
    # Neo4j stats
    try:
        neo4j_stats = neo4j_service.get_case_statistics()
        stats["neo4j"] = neo4j_stats
    except Exception as e:
        stats["neo4j"] = {"error": "Not available"}
    
    return stats

@app.get("/api/case/{case_id}/similar")
def get_similar_cases(case_id: str):
    """Find similar cases using Neo4j graph traversal"""
    try:
        similar = neo4j_service.find_similar_cases(case_id, limit=5)
        return {"similar_cases": similar}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Neo4j query failed: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
